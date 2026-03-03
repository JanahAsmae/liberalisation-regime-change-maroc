"""
Extraction des données IPC depuis les fichiers DOCX (janvier de chaque année).
- Prend le premier tableau (supposé être celui des divisions de produits).
- Colonne 0 : libellé de la division.
- Colonne 2 : valeur du mois courant (janvier).
- L'année est extraite du nom du fichier (ex: ipc_2013_01.docx → 2013).
- Mois = 'janvier' (constant pour ces fichiers).
"""

import os
import re
import pandas as pd
from docx import Document

def extraire_annee_depuis_nom(fichier):
    """Extrait l'année du nom de fichier (ex: ipc_2013_01.docx → 2013)."""
    match = re.search(r'ipc_(\d{4})_\d{2}\.docx', fichier)
    if match:
        return int(match.group(1))
    else:
        raise ValueError(f"Impossible d'extraire l'année de {fichier}")

def parse_float(val):
    if val is None:
        return None
    val = str(val).strip().replace(',', '.')
    try:
        return float(val)
    except ValueError:
        return None

def extraire_tableau_divisions(doc):
    """
    Parcourt les tableaux du document et retourne celui qui contient
    "Divisions de produits" (ou le premier s'il n'y en a qu'un).
    Retourne la liste des lignes (chaque ligne = liste de cellules text).
    """
    for table in doc.tables:
        # Vérifier si la première cellule contient "Divisions de produits"
        if table.rows and table.rows[0].cells:
            first_cell = table.rows[0].cells[0].text.strip()
            if "Divisions de produits" in first_cell:
                # Convertir en liste de listes
                data = []
                for row in table.rows:
                    data.append([cell.text.strip() for cell in row.cells])
                return data
    # Si non trouvé, prendre le premier tableau
    if doc.tables:
        table = doc.tables[0]
        data = []
        for row in table.rows:
            data.append([cell.text.strip() for cell in row.cells])
        return data
    return None

def main():
    input_dir = "data/raw/maroc/inflation/IPC/docx/structure3"
    output_dir = "data/interim/maroc/inflation"
    os.makedirs(output_dir, exist_ok=True)

    fichiers = [f for f in os.listdir(input_dir) if f.endswith('.docx')]
    for fichier in sorted(fichiers):
        chemin = os.path.join(input_dir, fichier)
        print(f"Traitement : {fichier}")
        try:
            annee = extraire_annee_depuis_nom(fichier)
        except ValueError as e:
            print(f"  {e} → ignoré")
            continue

        doc = Document(chemin)
        data_table = extraire_tableau_divisions(doc)
        if not data_table:
            print("  Aucun tableau trouvé")
            continue

        # la ligne d'en-tête est la ligne 1 (index 1)
        # et la colonne du mois courant est l'index 2 (0-based)
        if len(data_table) < 3:
            print("  Tableau trop petit")
            continue

        # Vérification rapide : la cellule (1,2) doit contenir un mois (ex: jan.2013)
        # On extrait les lignes de données à partir de l'index 2
        records = []
        for row in data_table[2:]:
            if len(row) < 3:
                continue
            label = row[0].strip()
            if not label or label == "Ensemble" and len(records) > 0:  # on garde Ensemble
                pass
            valeur = parse_float(row[2])
            if valeur is not None:
                records.append({
                    "Division par produits": label,
                    "ipc": valeur,
                    "annee": annee,
                    "mois": "janvier"
                })
            # On continue pour inclure "Ensemble" même si c'est en dernière ligne

        if records:
            df = pd.DataFrame(records)
            # Supprimer doublons éventuels
            df = df.drop_duplicates(subset=["Division par produits", "annee", "mois"])
            # Trier selon l'ordre naturel (optionnel)
            df = df.sort_values("Division par produits").reset_index(drop=True)
            nom_csv = fichier.replace('.docx', '.csv')
            chemin_csv = os.path.join(output_dir, nom_csv)
            df.to_csv(chemin_csv, index=False, encoding='utf-8-sig')
            print(f"  -> {len(df)} enregistrements -> {nom_csv}")
        else:
            print("  Aucune donnée extraite")

if __name__ == "__main__":
    main()