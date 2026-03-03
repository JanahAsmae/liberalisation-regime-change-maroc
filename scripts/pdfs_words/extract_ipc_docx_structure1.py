"""
Extraction spécifique pour les fichiers DOCX ayant la structure suivante :
- Premier tableau de la page 2
- Ligne d'en-tête avec deux blocs : "Indices mensuels" et "Indices du mois de janvier"
- Colonnes : [label, déc.2012, jan.2013, var.%, 2012, 2013, var.%]
On extrait la colonne du mois courant (jan.2013, soit la troisième colonne, indice 2).
"""

import os
import re
import pandas as pd
from docx import Document

MOIS_FR = {
    "janvier": 1, "février": 2, "mars": 3, "avril": 4,
    "mai": 5, "juin": 6, "juillet": 7, "août": 8,
    "septembre": 9, "octobre": 10, "novembre": 11, "décembre": 12,
}

def parse_month_year(text):
    """Extrait (mois_num, annee, mois_nom) d'une chaîne comme 'janv. 2013'."""
    if not text:
        return None, None, None
    text = text.strip().lower()
    abbr = {
        "janv.": "janvier", "févr.": "février", "mars": "mars", "avr.": "avril",
        "mai": "mai", "juin": "juin", "juil.": "juillet", "août": "août",
        "sept.": "septembre", "oct.": "octobre", "nov.": "novembre", "déc.": "décembre"
    }
    for a, n in abbr.items():
        if a in text:
            text = text.replace(a, n)
    for nom, num in MOIS_FR.items():
        if nom in text:
            match = re.search(r"(\d{4})", text)
            if match:
                return num, int(match.group(1)), nom
    return None, None, None

def parse_float(val):
    if val is None:
        return None
    val = str(val).strip().replace(",", ".")
    try:
        return float(val)
    except ValueError:
        return None

def extract_current_month_from_table(table):
    """
    Extrait les données du tableau en prenant la colonne du mois courant (indice 2).
    La ligne d'en-tête est la deuxième (index 1), les données commencent à l'index 2.
    """
    data = []
    for row in table.rows:
        row_data = [cell.text.strip() for cell in row.cells]
        data.append(row_data)

    if len(data) < 3:
        return []  # pas assez de lignes

    # La ligne d'en-tête des mois est la deuxième (index 1)
    header_row = data[1]
    # On s'assure que la colonne 2 contient bien un mois
    mois_num, annee_cur, mois_cur_nom = parse_month_year(header_row[2])
    if mois_cur_nom is None:
        # Si la détection échoue, on utilise la colonne 2 par défaut (peut-être en-tête différent)
        # On peut essayer de chercher un mois ailleurs, mais on force la colonne 2
        # On extrait l'année et le mois à partir du nom de la colonne si possible
        # Sinon, on utilise l'année du fichier ? Mais c'est risqué.
        # Pour l'instant, on abandonne.
        print("  Attention : la colonne 2 ne contient pas un mois reconnu.")
        return []

    records = []
    for row in data[2:]:  # à partir de la troisième ligne
        if len(row) < 3:
            continue
        label = row[0] if len(row) > 0 else ""
        if not label:
            continue
        val_cur = parse_float(row[2])  # valeur du mois courant
        if val_cur is not None:
            records.append({
                "Division par produits": label,
                "ipc": val_cur,
                "annee": annee_cur,
                "mois": mois_cur_nom
            })
    return records

def main():
    input_dir = "data/raw/maroc/inflation/IPC/docx/structure1"   # à ajuster
    output_dir = "data/interim/maroc/inflation"
    os.makedirs(output_dir, exist_ok=True)

    docx_files = [f for f in os.listdir(input_dir) if f.endswith(".docx")]
    if not docx_files:
        print("Aucun fichier DOCX trouvé dans", input_dir)
        return

    for fname in docx_files:
        filepath = os.path.join(input_dir, fname)
        print(f"Traitement : {fname}")
        doc = Document(filepath)

        records = []
        if doc.tables:
            table = doc.tables[0]  # premier tableau
            records.extend(extract_current_month_from_table(table))
        else:
            print("  -> Aucun tableau trouvé")

        if records:
            df = pd.DataFrame(records)
            df = df.drop_duplicates(subset=["Division par produits", "annee", "mois"])
            # Trier (optionnel)
            df["mois_num"] = df["mois"].map(MOIS_FR)
            df = df.sort_values(["annee", "mois_num", "Division par produits"]).drop(columns=["mois_num"])
            csv_name = os.path.splitext(fname)[0] + ".csv"
            csv_path = os.path.join(output_dir, csv_name)
            df.to_csv(csv_path, index=False, encoding="utf-8-sig")
            print(f"  -> {len(df)} enregistrements sauvegardés dans {csv_name}")
        else:
            print("  -> Aucune donnée extraite")

if __name__ == "__main__":
    main()