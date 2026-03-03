"""
Extraction des données IPC depuis des fichiers .doc (format Word ancien).
Utilise une conversion temporaire en .docx via LibreOffice.
"""

import os
import re
import subprocess
import tempfile
import shutil
import pandas as pd
from docx import Document

MOIS_FR = {
    "janvier": 1, "février": 2, "mars": 3, "avril": 4,
    "mai": 5, "juin": 6, "juillet": 7, "août": 8,
    "septembre": 9, "octobre": 10, "novembre": 11, "décembre": 12,
}

def parse_month_year(text):
    """Extrait (mois_num, annee, mois_nom) d'une chaîne comme 'janv. 2011'."""
    if not text:
        return None, None, None
    text = text.strip().lower()
    abbr = {
        "janv.": "janvier", "févr.": "février", "mars": "mars", "avr.": "avril",
        "mai": "mai", "juin": "juin", "juil.": "juillet", "août": "août",
        "sept.": "septembre", "oct.": "octobre", "nov.": "novembre", "déc.": "décembre"
    }
    for a, n in abbr.items():
        if a in text:
            text = text.replace(a, n)
    for nom, num in MOIS_FR.items():
        if nom in text:
            match = re.search(r"(\d{4})", text)
            if match:
                return num, int(match.group(1)), nom
    return None, None, None

def parse_float(val):
    if val is None:
        return None
    val = str(val).strip().replace(",", ".")
    try:
        return float(val)
    except ValueError:
        return None

def extract_from_docx_table(table):
    """
    Extrait les données d'un tableau DOCX (objet python-docx).
    La deuxième ligne contient les en-têtes des mois,
    les lignes suivantes contiennent les données.
    Il faut sélectionne la colonne du mois le plus récent.
    """
    data = []
    for row in table.rows:
        row_data = [cell.text.strip() for cell in row.cells]
        data.append(row_data)

    if len(data) < 3:
        return []

    header_row = data[1]
    mois_indices = []
    for i, cell in enumerate(header_row):
        mois_num, annee, mois_nom = parse_month_year(cell)
        if mois_nom:
            mois_indices.append((i, annee, mois_num, mois_nom))

    if not mois_indices:
        return []

    mois_indices.sort(key=lambda x: (x[1], x[2]), reverse=True)
    col_index_cur, annee_cur, _, mois_cur_nom = mois_indices[0]

    records = []
    for row in data[2:]:
        if len(row) <= col_index_cur:
            continue
        label = row[0] if len(row) > 0 else ""
        if not label:
            continue
        val_cur = parse_float(row[col_index_cur])
        if val_cur is not None:
            records.append({
                "Division par produits": label,
                "ipc": val_cur,
                "annee": annee_cur,
                "mois": mois_cur_nom
            })
    return records

def convert_doc_to_docx(doc_path):
    """
    Convertit un fichier .doc en .docx temporaire en utilisant LibreOffice.
    Retourne le chemin du fichier .docx converti, ou None si échec.
    """
    soffice_path = r"C:\Program Files\LibreOffice\program\soffice.exe"
    if not os.path.isfile(soffice_path):
        print(f"  ERREUR : LibreOffice introuvable à {soffice_path}")
        return None

    # Vérifier que le fichier source existe
    if not os.path.isfile(doc_path):
        print(f"  ERREUR : Fichier source introuvable : {doc_path}")
        return None

    doc_path_abs = os.path.abspath(doc_path)
    print(f"  Fichier source : {doc_path_abs} (taille : {os.path.getsize(doc_path_abs)} octets)")

    # Créer un répertoire temporaire pour la sortie
    with tempfile.TemporaryDirectory() as tmpdir:
        cmd = [
            soffice_path,
            "--headless",
            "--convert-to", "docx",
            "--outdir", tmpdir,
            doc_path_abs
        ]
        # Afficher la commande exacte pour diagnostic
        cmd_str = " ".join(f'"{c}"' for c in cmd)
        print(f"  Commande exécutée : {cmd_str}")

        try:
            result = subprocess.run(cmd, check=True, capture_output=True, text=True)
            print(f"  Sortie de LibreOffice : {result.stdout.strip()}")
        except subprocess.CalledProcessError as e:
            print(f"  Erreur de conversion (code {e.returncode}) : {e.stderr}")
            return None

        # Chercher le fichier converti
        base = os.path.basename(doc_path_abs).replace(".doc", ".docx")
        converted = os.path.join(tmpdir, base)
        if os.path.exists(converted):
            # Copier vers un fichier temporaire persistant
            final_path = os.path.join(tempfile.gettempdir(), base)
            shutil.copy2(converted, final_path)
            print(f"  Conversion réussie : {final_path}")
            return final_path
        else:
            print("  Fichier converti introuvable dans le répertoire temporaire.")
            return None

def main():
    input_dir = "data/raw/maroc/inflation/IPC/doc/structure1"  # à adapter
    output_dir = "data/interim/maroc/inflation"
    os.makedirs(output_dir, exist_ok=True)

    doc_files = [f for f in os.listdir(input_dir) if f.endswith(".doc")]

    if not doc_files:
        print("Aucun fichier .doc trouvé dans", input_dir)
        return

    for fname in doc_files:
        filepath = os.path.join(input_dir, fname)
        print(f"\nTraitement : {fname}")

        docx_path = convert_doc_to_docx(filepath)
        if docx_path is None:
            print("  -> Conversion échouée, fichier ignoré.")
            continue

        try:
            doc = Document(docx_path)
            records = []
            if doc.tables:
                table = doc.tables[0]
                records.extend(extract_from_docx_table(table))
            else:
                print("  -> Aucun tableau trouvé dans le document converti")

            if records:
                df = pd.DataFrame(records)
                df = df.drop_duplicates(subset=["Division par produits", "annee", "mois"])
                df["mois_num"] = df["mois"].map(MOIS_FR)
                df = df.sort_values(["annee", "mois_num", "Division par produits"]).drop(columns=["mois_num"])
                csv_name = os.path.splitext(fname)[0] + ".csv"
                csv_path = os.path.join(output_dir, csv_name)
                df.to_csv(csv_path, index=False, encoding="utf-8-sig")
                print(f"  -> {len(df)} enregistrements sauvegardés dans {csv_name}")
            else:
                print("  -> Aucune donnée extraite")

        finally:
            if os.path.exists(docx_path):
                os.unlink(docx_path)

if __name__ == "__main__":
    main()