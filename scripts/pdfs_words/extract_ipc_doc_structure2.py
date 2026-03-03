"""
Script spécifique pour extraire les données du fichier ipc_2012_01.doc.
- Conversion temporaire en .docx via LibreOffice.
- Extraction du premier tableau (divisions de produits).
- Colonne 0 : libellé, colonne 2 : valeur du mois courant (janvier).
- Année extraite du nom de fichier (2012), mois = 'janvier'.
"""

import os
import re
import subprocess
import tempfile
import shutil
import pandas as pd
from docx import Document

# ----------------------------------------------------------------------
# Configuration
# ----------------------------------------------------------------------
INPUT_FILE = "data/raw/maroc/inflation/IPC/doc/structure2/ipc_2012_01.doc"
OUTPUT_DIR = "data/interim/maroc/inflation"
SOFFICE_PATH = r"C:\Program Files\LibreOffice\program\soffice.exe"

# ----------------------------------------------------------------------
# Fonctions
# ----------------------------------------------------------------------
def parse_float(val):
    """Convertit une chaîne avec virgule en float."""
    if val is None:
        return None
    val = str(val).strip().replace(',', '.')
    try:
        return float(val)
    except ValueError:
        return None

def extraire_annee_depuis_nom(fichier):
    """Extrait l'année du nom de fichier (ex: ipc_2012_01.doc → 2012)."""
    match = re.search(r'ipc_(\d{4})_\d{2}\.doc', os.path.basename(fichier))
    if match:
        return int(match.group(1))
    else:
        raise ValueError(f"Impossible d'extraire l'année de {fichier}")

def convert_doc_to_docx(doc_path):
    """
    Convertit un fichier .doc en .docx temporaire via LibreOffice.
    Retourne le chemin du .docx converti, ou None en cas d'échec.
    """
    if not os.path.isfile(SOFFICE_PATH):
        print(f"ERREUR : LibreOffice introuvable à {SOFFICE_PATH}")
        return None

    if not os.path.isfile(doc_path):
        print(f"ERREUR : Fichier source introuvable : {doc_path}")
        return None

    doc_path_abs = os.path.abspath(doc_path)
    print(f"Fichier source : {doc_path_abs} (taille : {os.path.getsize(doc_path_abs)} octets)")

    # Créer un répertoire temporaire pour la sortie
    with tempfile.TemporaryDirectory() as tmpdir:
        cmd = [
            SOFFICE_PATH,
            "--headless",
            "--convert-to", "docx",
            "--outdir", tmpdir,
            doc_path_abs
        ]
        cmd_str = " ".join(f'"{c}"' for c in cmd)
        print(f"Commande exécutée : {cmd_str}")

        try:
            result = subprocess.run(cmd, check=True, capture_output=True, text=True)
            print(f"Sortie de LibreOffice : {result.stdout.strip()}")
        except subprocess.CalledProcessError as e:
            print(f"Erreur de conversion (code {e.returncode}) : {e.stderr}")
            return None

        base = os.path.basename(doc_path_abs).replace(".doc", ".docx")
        converted = os.path.join(tmpdir, base)
        if os.path.exists(converted):
            final_path = os.path.join(tempfile.gettempdir(), base)
            shutil.copy2(converted, final_path)
            print(f"Conversion réussie : {final_path}")
            return final_path
        else:
            print("Fichier converti introuvable dans le répertoire temporaire.")
            return None

def extraire_tableau_divisions(doc):
    """
    Parcourt les tableaux du document et retourne celui qui contient
    "Divisions de produits" (ou le premier sinon).
    Retourne la liste des lignes (chaque ligne = liste de cellules text).
    """
    for table in doc.tables:
        if table.rows and table.rows[0].cells:
            first_cell = table.rows[0].cells[0].text.strip()
            if "Divisions de produits" in first_cell:
                data = []
                for row in table.rows:
                    data.append([cell.text.strip() for cell in row.cells])
                return data
    if doc.tables:
        table = doc.tables[0]
        data = []
        for row in table.rows:
            data.append([cell.text.strip() for cell in row.cells])
        return data
    return None

def main():
    # Créer le dossier de sortie
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    # Vérifier que le fichier source existe
    if not os.path.isfile(INPUT_FILE):
        print(f"Fichier source introuvable : {INPUT_FILE}")
        return

    # Extraire l'année du nom
    try:
        annee = extraire_annee_depuis_nom(INPUT_FILE)
    except ValueError as e:
        print(e)
        return

    print(f"Traitement de {os.path.basename(INPUT_FILE)} (année {annee})")

    # Conversion en .docx
    docx_path = convert_doc_to_docx(INPUT_FILE)
    if docx_path is None:
        print("Conversion échouée, arrêt.")
        return

    try:
        # Charger le document converti
        doc = Document(docx_path)
        data_table = extraire_tableau_divisions(doc)
        if not data_table:
            print("Aucun tableau trouvé dans le document converti.")
            return

        if len(data_table) < 3:
            print("Tableau trop petit.")
            return

        # Parcourir les lignes de données à partir de l'index 2
        records = []
        for row in data_table[2:]:
            if len(row) < 3:
                continue
            label = row[0].strip()
            if not label:
                continue
            valeur = parse_float(row[2])   # colonne du mois courant (janvier)
            if valeur is not None:
                records.append({
                    "Division par produits": label,
                    "ipc": valeur,
                    "annee": annee,
                    "mois": "janvier"
                })

        if not records:
            print("Aucune donnée extraite.")
            return

        # Créer le DataFrame et nettoyer
        df = pd.DataFrame(records)
        df = df.drop_duplicates(subset=["Division par produits", "annee", "mois"])
        df = df.sort_values("Division par produits").reset_index(drop=True)

        # Sauvegarder en CSV
        nom_csv = os.path.basename(INPUT_FILE).replace(".doc", ".csv")
        chemin_csv = os.path.join(OUTPUT_DIR, nom_csv)
        df.to_csv(chemin_csv, index=False, encoding='utf-8-sig')
        print(f"Fichier CSV sauvegardé : {chemin_csv}")
        print(f"Nombre d'enregistrements : {len(df)}")

    finally:
        # Nettoyer le fichier .docx temporaire
        if os.path.exists(docx_path):
            os.unlink(docx_path)

if __name__ == "__main__":
    main()