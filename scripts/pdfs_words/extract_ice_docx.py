# """
# Extraction des données ICE (Indices du Commerce Extérieur) depuis les fichiers DOCX.
# - Prend le premier tableau de chaque document.
# - Extrait les données pour les importations et exportations.
# - Colonnes : Division par produits, ice, annee, trimestre, type.
# """

# import os
# import re
# import pandas as pd
# from docx import Document

# def extraire_annee_trimestre(nom_fichier):
#     """
#     Extrait l'année et le trimestre depuis un nom comme 'ice_t1_2025_fr.docx'.
#     Retourne (annee, trimestre) où trimestre est 't1', 't2', 't3', 't4'.
#     """
#     pattern = r'ice_(t[1-4])_(\d{4})_fr\.docx'
#     match = re.search(pattern, nom_fichier)
#     if match:
#         trimestre = match.group(1)  # 't1', 't2', etc.
#         annee = int(match.group(2))
#         return annee, trimestre
#     else:
#         raise ValueError(f"Nom de fichier invalide : {nom_fichier}")

# def parse_float(val):
#     """Convertit une chaîne avec virgule en float."""
#     if val is None:
#         return None
#     val = str(val).strip().replace(',', '.')
#     try:
#         return float(val)
#     except ValueError:
#         return None

# def extraire_tableau_ice(table):
#     """
#     Parcourt le tableau et extrait les lignes pour IMPORTATIONS et EXPORTATIONS.
#     Retourne une liste de dictionnaires avec les clés : division, ice, type.
#     """
#     data = []
#     for row in table.rows:
#         row_data = [cell.text.strip() for cell in row.cells]
#         data.append(row_data)

#     if len(data) < 3:
#         return []

#     records = []
#     current_type = None  # 'import' ou 'export'

#     for row in data:
#         if len(row) < 3:
#             continue
#         label = row[0].strip() if row[0] else ""
#         if not label:
#             continue

#         # Détection des en-têtes de section
#         if "IMPORTATIONS" in label.upper() and "EXPORTATIONS" not in label.upper():
#             current_type = "import"
#             continue
#         elif "EXPORTATIONS" in label.upper():
#             current_type = "export"
#             continue

#         # Si on est dans une section et que la ligne a un libellé de produit
#         if current_type and label and not label.startswith(("IMPORTATIONS", "EXPORTATIONS")):
#             # La valeur de l'ICE est dans la colonne 2 (celle du trimestre courant)
#             ice_val = parse_float(row[2])
#             if ice_val is not None:
#                 records.append({
#                     "division": label,
#                     "ice": ice_val,
#                     "type": current_type
#                 })
#     return records

# def main():
#     input_dir = "data/raw/maroc/import et export(ice)/docx"
#     output_dir = "data/interim/maroc/ICE"
#     os.makedirs(output_dir, exist_ok=True)

#     fichiers = [f for f in os.listdir(input_dir) if f.endswith('.docx') and f.startswith('ice_t')]
#     if not fichiers:
#         print("Aucun fichier correspondant trouvé.")
#         return

#     for fichier in sorted(fichiers):
#         chemin = os.path.join(input_dir, fichier)
#         print(f"Traitement : {fichier}")

#         try:
#             annee, trimestre = extraire_annee_trimestre(fichier)
#         except ValueError as e:
#             print(f"  {e} → ignoré")
#             continue

#         doc = Document(chemin)
#         if not doc.tables:
#             print("  Aucun tableau trouvé.")
#             continue

#         # Premier tableau
#         table = doc.tables[0]
#         records = extraire_tableau_ice(table)

#         if not records:
#             print("  Aucune donnée extraite.")
#             continue

#         # Ajouter l'année et le trimestre
#         for r in records:
#             r["annee"] = annee
#             r["trimestre"] = trimestre

#         df = pd.DataFrame(records)
#         # Réordonner les colonnes
#         df = df[["division", "ice", "annee", "trimestre", "type"]]
#         # Supprimer doublons éventuels
#         df = df.drop_duplicates()
#         # Trier par type puis division (optionnel)
#         df = df.sort_values(["type", "division"]).reset_index(drop=True)

#         nom_csv = fichier.replace('.docx', '.csv')
#         chemin_csv = os.path.join(output_dir, nom_csv)
#         df.to_csv(chemin_csv, index=False, encoding='utf-8-sig')
#         print(f"  -> {len(df)} enregistrements sauvegardés dans {nom_csv}")

# if __name__ == "__main__":
#     main()

"""
Extraction des données ICE (Indices du Commerce Extérieur) depuis les fichiers DOCX.
- Prend le premier tableau de chaque document.
- Extrait les données pour les importations et exportations, y compris les totaux.
- Colonnes : division, ice, annee, trimestre, type.
"""

import os
import re
import pandas as pd
from docx import Document

def extraire_annee_trimestre(nom_fichier):
    """
    Extrait l'année et le trimestre depuis un nom comme 'ice_t1_2025_fr.docx'.
    Retourne (annee, trimestre) où trimestre est 't1', 't2', 't3', 't4'.
    """
    pattern = r'ice_(t[1-4])_(\d{4})_fr\.docx'
    match = re.search(pattern, nom_fichier)
    if match:
        trimestre = match.group(1)  # 't1', 't2', etc.
        annee = int(match.group(2))
        return annee, trimestre
    else:
        raise ValueError(f"Nom de fichier invalide : {nom_fichier}")

def parse_float(val):
    """Convertit une chaîne avec virgule en float."""
    if val is None:
        return None
    val = str(val).strip().replace(',', '.')
    try:
        return float(val)
    except ValueError:
        return None

def extraire_tableau_ice(table):
    """
    Parcourt le tableau et extrait les lignes pour IMPORTATIONS et EXPORTATIONS.
    Retourne une liste de dictionnaires avec les clés : division, ice, type.
    Inclut les lignes "IMPORTATIONS" et "EXPORTATIONS" elles-mêmes.
    """
    data = []
    for row in table.rows:
        row_data = [cell.text.strip() for cell in row.cells]
        data.append(row_data)

    if len(data) < 3:
        return []

    records = []
    current_type = None  # 'import' ou 'export'

    for row in data:
        if len(row) < 3:
            continue
        label = row[0].strip() if row[0] else ""
        if not label:
            continue

        # Détection des en-têtes de section et extraction de leur valeur
        if "IMPORTATIONS" in label.upper() and "EXPORTATIONS" not in label.upper():
            current_type = "import"
            val = parse_float(row[2])
            if val is not None:
                records.append({
                    "division": label,
                    "ice": val,
                    "type": current_type
                })
            continue
        elif "EXPORTATIONS" in label.upper():
            current_type = "export"
            val = parse_float(row[2])
            if val is not None:
                records.append({
                    "division": label,
                    "ice": val,
                    "type": current_type
                })
            continue

        # Si on est dans une section et que la ligne a un libellé de produit
        if current_type and label and not label.startswith(("IMPORTATIONS", "EXPORTATIONS")):
            ice_val = parse_float(row[2])
            if ice_val is not None:
                records.append({
                    "division": label,
                    "ice": ice_val,
                    "type": current_type
                })
    return records

def main():
    input_dir = "data/raw/maroc/import et export(ice)/docx" 
    output_dir = "data/interim/maroc/ICE"
    os.makedirs(output_dir, exist_ok=True)

    fichiers = [f for f in os.listdir(input_dir) if f.endswith('.docx') and f.startswith('ice_t')]
    if not fichiers:
        print("Aucun fichier correspondant trouvé.")
        return

    for fichier in sorted(fichiers):
        chemin = os.path.join(input_dir, fichier)
        print(f"Traitement : {fichier}")

        try:
            annee, trimestre = extraire_annee_trimestre(fichier)
        except ValueError as e:
            print(f"  {e} → ignoré")
            continue

        doc = Document(chemin)
        if not doc.tables:
            print("  Aucun tableau trouvé.")
            continue

        # Premier tableau
        table = doc.tables[0]
        records = extraire_tableau_ice(table)

        if not records:
            print("  Aucune donnée extraite.")
            continue

        # Ajouter l'année et le trimestre
        for r in records:
            r["annee"] = annee
            r["trimestre"] = trimestre

        df = pd.DataFrame(records)
        # Réordonner les colonnes
        df = df[["division", "ice", "annee", "trimestre", "type"]]
        # Supprimer doublons éventuels
        df = df.drop_duplicates()
        # Trier par type puis division (optionnel)
        df = df.sort_values(["type", "division"]).reset_index(drop=True)

        nom_csv = fichier.replace('.docx', '.csv')
        chemin_csv = os.path.join(output_dir, nom_csv)
        df.to_csv(chemin_csv, index=False, encoding='utf-8-sig')
        print(f"  -> {len(df)} enregistrements sauvegardés dans {nom_csv}")

if __name__ == "__main__":
    main()