"""
Extraction des données IPC depuis des fichiers DOCX (uniquement le premier tableau).
Structure attendue : tableau avec colonnes (label, mois_ref, mois_cur, ...)
Seul le mois courant (le plus récent) est conservé.
"""

import os
import re
import pandas as pd
from docx import Document

MOIS_FR = {
    "janvier": 1, "février": 2, "mars": 3, "avril": 4,
    "mai": 5, "juin": 6, "juillet": 7, "août": 8,
    "septembre": 9, "octobre": 10, "novembre": 11, "décembre": 12,
}

def parse_month_year(text):
    """Extrait (mois_num, annee, mois_nom) d'une chaîne comme 'janv. 2011'."""
    if not text:
        return None, None, None
    text = text.strip().lower()
    # Remplacer les abréviations courantes
    abbr = {
        "janv.": "janvier", "févr.": "février", "mars": "mars", "avr.": "avril",
        "mai": "mai", "juin": "juin", "juil.": "juillet", "août": "août",
        "sept.": "septembre", "oct.": "octobre", "nov.": "novembre", "déc.": "décembre"
    }
    for a, n in abbr.items():
        if a in text:
            text = text.replace(a, n)
    for nom, num in MOIS_FR.items():
        if nom in text:
            match = re.search(r"(\d{4})", text)
            if match:
                return num, int(match.group(1)), nom
    return None, None, None

def parse_float(val):
    if val is None:
        return None
    val = str(val).strip().replace(",", ".")
    try:
        return float(val)
    except ValueError:
        return None

def extract_from_docx_table(table):
    """
    Extrait les données d'un tableau DOCX.
    La deuxième ligne contient les en-têtes des mois,
    les lignes suivantes contiennent les données.
    Sélectionne la colonne du mois le plus récent (mois courant).
    """
    data = []
    for row in table.rows:
        row_data = [cell.text.strip() for cell in row.cells]
        data.append(row_data)

    if len(data) < 3:
        return []  # pas assez de lignes

    # Ligne d'en-tête des mois (index 1)
    header_row = data[1]

    # Collecter tous les indices de colonnes contenant un mois
    mois_indices = []
    for i, cell in enumerate(header_row):
        mois_num, annee, mois_nom = parse_month_year(cell)
        if mois_nom:
            mois_indices.append((i, annee, mois_num, mois_nom))

    if not mois_indices:
        return []  # aucun mois trouvé

    # Trier par année décroissante puis par mois décroissant
    mois_indices.sort(key=lambda x: (x[1], x[2]), reverse=True)
    col_index_cur, annee_cur, _, mois_cur_nom = mois_indices[0]

    records = []
    for row in data[2:]:  # à partir de la troisième ligne
        if len(row) <= col_index_cur:
            continue
        label = row[0] if len(row) > 0 else ""
        if not label:
            continue
        val_cur = parse_float(row[col_index_cur])
        if val_cur is not None:
            records.append({
                "Division par produits": label,
                "ipc": val_cur,
                "annee": annee_cur,
                "mois": mois_cur_nom
            })
    return records

def main():
    # À adapter selon l'emplacement de vos fichiers
    input_dir = "data/raw/maroc/inflation/IPC/docx/structure2"
    output_dir = "data/interim/maroc/inflation"
    os.makedirs(output_dir, exist_ok=True)

    docx_files = [f for f in os.listdir(input_dir) if f.endswith(".docx")]
    if not docx_files:
        print("Aucun fichier DOCX trouvé dans", input_dir)
        return

    for fname in docx_files:
        filepath = os.path.join(input_dir, fname)
        print(f"Traitement : {fname}")
        doc = Document(filepath)

        records = []
        # Ne prendre que le premier tableau du document
        if doc.tables:
            table = doc.tables[0]
            records.extend(extract_from_docx_table(table))
        else:
            print("  -> Aucun tableau trouvé dans le document")

        if records:
            df = pd.DataFrame(records)
            df = df.drop_duplicates(subset=["Division par produits", "annee", "mois"])
            # Trier par année, mois (numéro) et division
            df["mois_num"] = df["mois"].map(MOIS_FR)
            df = df.sort_values(["annee", "mois_num", "Division par produits"]).drop(columns=["mois_num"])
            csv_name = os.path.splitext(fname)[0] + ".csv"
            csv_path = os.path.join(output_dir, csv_name)
            df.to_csv(csv_path, index=False, encoding="utf-8-sig")
            print(f"  -> {len(df)} enregistrements sauvegardés dans {csv_name}")
        else:
            print("  -> Aucune donnée extraite (tableau vide ou mal formaté)")

if __name__ == "__main__":
    main()